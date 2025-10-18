#!/usr/bin/env python3
"""
docling_extract.py

Batch-convert .pdf and .docx files in ./documents to text using local extraction,
saving .txt files to ./output.

This replaces the LLMWhisperer API-based extraction with local extraction using:
- PyMuPDF (fitz) for PDF files
- python-docx for DOCX files
- Tesseract OCR for scanned PDFs (optional)

USAGE:
  python3 docling_extract.py
  python3 docling_extract.py --in documents --out output
  python3 docling_extract.py --in documents --out output --ocr
  python3 docling_extract.py --in documents --out output --force-ocr

REQUIREMENTS:
  pip install pymupdf python-docx
  
  For OCR support (optional):
  pip install pytesseract pillow
  sudo apt-get install tesseract-ocr
"""

from __future__ import annotations
import argparse
import sys
from pathlib import Path
from typing import Optional

try:
    import fitz  # PyMuPDF
except ImportError:
    print("ERROR: PyMuPDF not installed. Run: pip install pymupdf", file=sys.stderr)
    sys.exit(2)

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(2)

# OCR support is optional
OCR_AVAILABLE = False
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    pass  # OCR will not be available


def ensure_output_dir(p: Path) -> None:
    """Create output directory if it doesn't exist."""
    p.mkdir(parents=True, exist_ok=True)


def has_text_layer(pdf_doc: fitz.Document) -> bool:
    """
    Check if a PDF has an embedded text layer.
    
    Args:
        pdf_doc: Opened PDF document
        
    Returns:
        True if PDF has extractable text, False otherwise
    """
    # Sample first few pages to check for text
    max_pages_to_check = min(3, len(pdf_doc))
    total_text_length = 0
    
    for page_num in range(max_pages_to_check):
        page = pdf_doc[page_num]
        text = page.get_text("text").strip()
        total_text_length += len(text)
    
    # If we have at least some text across the pages, consider it has text layer
    # Threshold: at least 100 characters across sampled pages
    return total_text_length > 100


def extract_text_with_ocr(pdf_doc: fitz.Document) -> str:
    """
    Extract text from PDF using OCR (for scanned documents).
    
    Args:
        pdf_doc: Opened PDF document
        
    Returns:
        OCR-extracted text content
    """
    if not OCR_AVAILABLE:
        return "[OCR NOT AVAILABLE] PDF appears to be scanned but pytesseract is not installed.\n" \
               "Install with: pip install pytesseract pillow\n" \
               "Also install: sudo apt-get install tesseract-ocr"
    
    text_parts = []
    
    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        
        # Render page to image at 300 DPI for good OCR quality
        mat = fitz.Matrix(300/72, 300/72)  # 300 DPI
        pix = page.get_pixmap(matrix=mat)
        
        # Convert to PIL Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        # Run OCR
        try:
            text = pytesseract.image_to_string(img)
            if text.strip():
                text_parts.append(text)
        except Exception as e:
            text_parts.append(f"[OCR ERROR on page {page_num + 1}: {e}]")
    
    return "\n".join(text_parts)


def extract_text_from_pdf(file_path: Path, use_ocr: bool = False, force_ocr: bool = False, auto_ocr: bool = True) -> str:
    """
    Extract text from a PDF file using PyMuPDF, with automatic OCR fallback.
    
    Priority 1.1: OCR Auto-Detection Enhancement
    - Automatically detects when a PDF has no text layer
    - Triggers OCR automatically if OCR dependencies are available
    - Provides clear feedback about OCR usage
    
    Args:
        file_path: Path to the PDF file
        use_ocr: Enable OCR fallback for scanned PDFs (default: False, deprecated - use auto_ocr)
        force_ocr: Force OCR even if text layer exists (default: False)
        auto_ocr: Automatically use OCR if no text layer detected (default: True, new in Priority 1.1)
        
    Returns:
        Extracted text content
    """
    doc = fitz.open(file_path)
    
    # If force_ocr, always use OCR
    if force_ocr:
        if OCR_AVAILABLE:
            print(f"  [OCR] Force OCR enabled for {file_path.name}")
            text = extract_text_with_ocr(doc)
        else:
            print(f"  [WARN] Force OCR requested but OCR not available", file=sys.stderr)
            text = extract_text_normally(doc)
    # Priority 1.1: Automatic OCR detection
    elif not has_text_layer(doc):
        # Automatically use OCR if available (unless disabled by auto_ocr=False)
        if auto_ocr and OCR_AVAILABLE:
            print(f"  [AUTO-OCR] PDF appears to be scanned, automatically using OCR for {file_path.name}")
            text = extract_text_with_ocr(doc)
        elif use_ocr and OCR_AVAILABLE:
            # Legacy path for backward compatibility
            print(f"  [OCR] No text layer detected, using OCR for {file_path.name}")
            text = extract_text_with_ocr(doc)
        elif not OCR_AVAILABLE:
            print(f"  [WARN] No text layer detected in {file_path.name}. OCR dependencies not installed.", file=sys.stderr)
            print(f"  [WARN] Install with: pip install pytesseract pillow && sudo apt-get install tesseract-ocr", file=sys.stderr)
            text = "[NO TEXT LAYER] This PDF appears to be scanned but OCR is not available.\n" \
                   "Install OCR: pip install pytesseract pillow && sudo apt-get install tesseract-ocr"
        else:
            # auto_ocr disabled and use_ocr not set
            print(f"  [WARN] No text layer detected in {file_path.name}. Use --ocr to enable OCR.", file=sys.stderr)
            text = "[NO TEXT LAYER] This PDF appears to be scanned. Re-run with --ocr flag or enable auto-OCR."
    else:
        # Normal text extraction (with page-level OCR fallback if enabled)
        text = extract_text_normally(doc, auto_ocr=auto_ocr)
    
    doc.close()
    return text


def extract_text_normally(pdf_doc: fitz.Document, auto_ocr: bool = True) -> str:
    """
    Extract text normally from PDF with embedded text layer.
    
    Patch 1: Page-Level OCR Fallback
    - If auto_ocr is enabled and a page has no text, attempts OCR on that page
    - Handles mixed-content PDFs (some pages scanned, some with text)
    
    Args:
        pdf_doc: Opened PDF document
        auto_ocr: Enable automatic OCR for blank pages (default: True)
        
    Returns:
        Extracted text content
    """
    text_parts = []
    
    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        # Use "text" mode for simple text extraction
        page_text = page.get_text("text").strip()
        
        # Patch 1: If page is blank and OCR available, try OCR on just this page
        if auto_ocr and OCR_AVAILABLE and not page_text:
            print(f"  [AUTO-OCR] Page {page_num+1} has no text, performing OCR")
            try:
                # Render page to image at 300 DPI
                mat = fitz.Matrix(300/72, 300/72)
                pix = page.get_pixmap(matrix=mat)
                
                # Convert to PIL Image
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Run OCR
                ocr_text = pytesseract.image_to_string(img)
                if ocr_text.strip():
                    text_parts.append(ocr_text)
            except Exception as e:
                print(f"  [WARN] OCR failed for page {page_num+1}: {e}")
                # Continue without this page's text
        elif page_text:
            text_parts.append(page_text)
    
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    doc = Document(file_path)
    text_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)


def unique_txt_path(dst: Path, source_path: Path) -> Path:
    """
    Generate a unique file path that accounts for concurrency and duplicate filenames.
    
    Patch 1: Non-atomic file naming fix
    - Uses folder hash to ensure unique output names even in parallel mode
    - Prevents race conditions when multiple files have the same base name
    
    Args:
        dst: Desired destination path (base name only)
        source_path: Original source file path (used for hash generation)
        
    Returns:
        Unique path that won't clash in parallel processing
    """
    # Generate a hash from the full source path to ensure uniqueness
    # This prevents conflicts when files with same name exist in different folders
    folder_hash = hex(abs(hash(str(source_path.parent.resolve()))) % (16**4))[2:].zfill(4)
    
    # Create output name with folder hash for uniqueness
    out_name = f"{dst.stem}_{folder_hash}{dst.suffix}"
    out_path = dst.parent / out_name
    
    # If still exists (very unlikely), append a number
    if not out_path.exists():
        return out_path
    
    i = 1
    while True:
        cand = dst.parent / f"{dst.stem}_{folder_hash}_{i}{dst.suffix}"
        if not cand.exists():
            return cand
        i += 1


def process_one(file_path: Path, out_dir: Path, use_ocr: bool = False, force_ocr: bool = False, auto_ocr: bool = True) -> None:
    """
    Process a single document file and extract text to output directory.
    
    Priority 1.1: OCR Auto-Detection Enhancement
    Patch 3: Skip non-extractable files
    
    Args:
        file_path: Path to the input document
        out_dir: Directory to save the extracted text
        use_ocr: Enable OCR fallback for scanned PDFs (legacy, prefer auto_ocr)
        force_ocr: Force OCR even if text layer exists
        auto_ocr: Automatically use OCR if no text layer detected (default: True)
    """
    print(f"[+] {file_path.name}")
    
    try:
        # Extract text based on file type
        if file_path.suffix.lower() == ".pdf":
            text = extract_text_from_pdf(file_path, use_ocr=use_ocr, force_ocr=force_ocr, auto_ocr=auto_ocr)
        elif file_path.suffix.lower() == ".docx":
            text = extract_text_from_docx(file_path)
        else:
            print(f"[!] Unsupported file type: {file_path.suffix}", file=sys.stderr)
            return
        
        # Patch 3: Skip files that could not be extracted (no text layer and no OCR)
        if text.startswith("[NO TEXT LAYER]") or text.startswith("[OCR NOT AVAILABLE]"):
            print(f"[!] Skipping {file_path.name} – no text and OCR unavailable", file=sys.stderr)
            print(f"    Install OCR: pip install pytesseract pillow && sudo apt-get install tesseract-ocr", file=sys.stderr)
            return
        
        # Save to output file with unique path (Patch 1: prevents race conditions)
        out_path = unique_txt_path(out_dir / (file_path.stem + ".txt"), file_path)
        out_path.write_text(text, encoding="utf-8")
        print(f"[✓] Saved -> {out_path}")
        
    except Exception as e:
        print(f"[x] Failed: {e}", file=sys.stderr)
        raise


def process_one_wrapper(args_tuple):
    """
    Wrapper for process_one to support multiprocessing.
    
    Priority 6.1: Parallel Processing Support
    - Enables parallel processing of multiple documents
    - Returns tuple of (success, filename, error_message)
    """
    file_path, out_dir, use_ocr, force_ocr, auto_ocr = args_tuple
    try:
        process_one(file_path=file_path, out_dir=out_dir, use_ocr=use_ocr, force_ocr=force_ocr, auto_ocr=auto_ocr)
        return (True, file_path.name, None)
    except Exception as e:
        return (False, file_path.name, str(e))


def main():
    ap = argparse.ArgumentParser(
        description="Extract text from PDF and DOCX files using local extraction"
    )
    ap.add_argument(
        "--in", 
        dest="in_dir", 
        default="documents", 
        help="Input folder (default: documents)"
    )
    ap.add_argument(
        "--out", 
        dest="out_dir", 
        default="output", 
        help="Output folder (default: output)"
    )
    ap.add_argument(
        "--ocr",
        action="store_true",
        help="Enable OCR fallback for scanned PDFs without text layer"
    )
    ap.add_argument(
        "--force-ocr",
        action="store_true",
        help="Force OCR for all PDFs, even if they have text layer"
    )
    ap.add_argument(
        "--no-auto-ocr",
        action="store_true",
        help="Disable automatic OCR for scanned PDFs (Priority 1.1: OCR Auto-Detection)"
    )
    ap.add_argument(
        "--jobs",
        type=int,
        default=1,
        help="Number of parallel jobs (default: 1 for sequential). Use -1 for CPU count."
    )
    args = ap.parse_args()

    in_dir = Path(args.in_dir)
    out_dir = Path(args.out_dir)
    ensure_output_dir(out_dir)

    if not in_dir.exists():
        print(f"ERROR: Input folder does not exist: {in_dir}", file=sys.stderr)
        sys.exit(2)

    # Priority 1.1: Auto-OCR is enabled by default (can be disabled with --no-auto-ocr)
    auto_ocr = not args.no_auto_ocr
    
    # Check OCR availability if requested or if auto_ocr is enabled
    if (args.ocr or args.force_ocr or auto_ocr) and not OCR_AVAILABLE:
        if auto_ocr:
            print("INFO: Auto-OCR is enabled but pytesseract not installed.", file=sys.stderr)
            print("Scanned PDFs will not be processed. Install with:", file=sys.stderr)
        else:
            print("WARNING: OCR requested but pytesseract not installed.", file=sys.stderr)
        print("  pip install pytesseract pillow", file=sys.stderr)
        print("  sudo apt-get install tesseract-ocr", file=sys.stderr)
        if not args.force_ocr and not args.ocr:
            print("Continuing with auto-OCR disabled...\n", file=sys.stderr)
        else:
            print("Continuing without OCR...\n", file=sys.stderr)

    # Find all PDF and DOCX files
    exts = {".pdf", ".docx"}
    files = sorted([p for p in in_dir.rglob("*") if p.is_file() and p.suffix.lower() in exts])
    
    if not files:
        print(f"WARNING: No .pdf or .docx files found in {in_dir}")
        return

    print(f"Found {len(files)} file(s) to process")
    mode_desc = "PyMuPDF for PDF, python-docx for DOCX"
    if args.force_ocr:
        mode_desc += " (Force OCR mode)"
    elif args.ocr:
        mode_desc += " (OCR fallback enabled)"
    elif auto_ocr and OCR_AVAILABLE:
        mode_desc += " (Auto-OCR enabled - Priority 1.1)"
    elif auto_ocr:
        mode_desc += " (Auto-OCR enabled but OCR not available)"
    
    # Priority 6.1: Determine number of jobs
    num_jobs = args.jobs
    if num_jobs == -1:
        import multiprocessing
        num_jobs = multiprocessing.cpu_count()
    
    print(f"Using local extraction ({mode_desc})")
    
    # Sequential processing (default)
    if num_jobs <= 1:
        for f in files:
            try:
                process_one(file_path=f, out_dir=out_dir, use_ocr=args.ocr, force_ocr=args.force_ocr, auto_ocr=auto_ocr)
            except Exception as e:
                print(f"[x] Failed on {f.name}: {e}", file=sys.stderr)
    
    # Parallel processing (Priority 6.1)
    else:
        import multiprocessing
        print(f"Processing {len(files)} file(s) with {num_jobs} parallel jobs...")
        
        # Prepare arguments for workers
        work_items = [(f, out_dir, args.ocr, args.force_ocr, auto_ocr) for f in files]
        
        # Process in parallel
        failed_files = []
        with multiprocessing.Pool(processes=num_jobs) as pool:
            results = pool.map(process_one_wrapper, work_items)
        
        # Report results
        successful = sum(1 for success, _, _ in results if success)
        for success, filename, error_msg in results:
            if not success:
                failed_files.append((filename, error_msg))
                print(f"[x] Failed on {filename}: {error_msg}", file=sys.stderr)
        
        print(f"\n✅ Completed: {successful}/{len(files)} files processed successfully")
        if failed_files:
            print(f"❌ Failed: {len(failed_files)} file(s) - see errors above")


if __name__ == "__main__":
    main()
